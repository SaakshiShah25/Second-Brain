"""
document_extract.py — Turns an uploaded client agreement (PDF, Word .docx,
or a photographed/scanned page) into structured deal terms.

Format handling (verified live before building this, not assumed):
  - PDF: hybrid extraction via PyMuPDF - tries each page's real embedded
    text layer first (fast, exact - this is what a Word->PDF export or any
    born-digital contract has), and only falls back to OCR (Tesseract) on
    pages that come back empty, which is what a scanned/photocopied signed
    page looks like (no text layer at all). This avoids OCR noise on the
    common case while still handling scans.
  - .docx: read via python-docx's structured XML (paragraphs + tables) -
    no OCR involved, most reliable of the three since Word never loses
    the text layer.
  - Images (.jpg/.png/etc, e.g. a phone photo of a printed page): OCR'd
    directly with Tesseract, same as card_scan.py's business-card path.
  - Legacy .doc (old binary Word format, pre-2007) is NOT supported -
    python-docx only reads .docx's XML format. Rare enough in practice
    (raise a clear error asking for .docx/PDF instead of silently
    mishandling it).

Only Tesseract (already required by card_scan.py) is needed as a system
binary; PyMuPDF and python-docx are pure Python wheels with no extra
system dependency.
"""

import json
from datetime import date
from io import BytesIO
from typing import Optional

import fitz  # PyMuPDF
import pytesseract
from dateutil.relativedelta import relativedelta
from docx import Document as DocxDocument
from PIL import Image

from llm_client import get_client, MODEL_NAME

# Below this many characters, a PDF page is treated as having no real text
# layer (e.g. a handful of stray characters from a stamp/watermark) and
# gets OCR'd instead of trusted as-is.
_MIN_NATIVE_TEXT_CHARS = 20

_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif"}


def _extension(filename: str) -> str:
    return (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()


def _ocr_image_bytes(image_bytes: bytes) -> str:
    image = Image.open(BytesIO(image_bytes))
    return pytesseract.image_to_string(image)


def _extract_pdf_text(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = []
    for page in doc:
        native = page.get_text().strip()
        if len(native) >= _MIN_NATIVE_TEXT_CHARS:
            pages_text.append(native)
        else:
            pix = page.get_pixmap(dpi=200)
            pages_text.append(pytesseract.image_to_string(Image.open(BytesIO(pix.tobytes("png")))).strip())
    return "\n\n".join(p for p in pages_text if p)


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatches on file extension. Raises ValueError for unsupported
    formats (notably legacy .doc) or if nothing readable was found."""
    ext = _extension(filename)
    if ext == "pdf":
        text = _extract_pdf_text(file_bytes)
    elif ext == "docx":
        text = _extract_docx_text(file_bytes)
    elif ext in _IMAGE_EXTENSIONS:
        text = _ocr_image_bytes(file_bytes)
    elif ext == "doc":
        raise ValueError("Legacy .doc files aren't supported - please save/export as .docx or PDF.")
    else:
        raise ValueError(f"Unsupported file type '.{ext}' - upload a PDF, .docx, or a clear photo of the document.")

    if not text or not text.strip():
        raise ValueError("Couldn't read any text from that document - try a clearer scan/photo, or a different file.")
    return text


_SYSTEM_PROMPT = """You are extracting structured deal terms from a client agreement/contract
document (the raw text may contain OCR noise, odd spacing, or line-break artifacts from a scan -
ignore that and focus on the actual content).

Return ONLY valid JSON (no markdown fences, no preamble) matching this exact schema:

{
  "client_company": "string - the CLIENT's company/organization name (not the provider's own company) - the party purchasing/receiving the service. Empty string if genuinely unclear.",
  "client_legal_name": "string - the client's full registered legal name if stated and different from client_company (e.g. 'Vertex Robotics Corp.' vs a shorter 'Vertex Robotics'), else empty string.",
  "provider_legal_name": "string - the provider/vendor's (the user's own company's) legal name as stated in the agreement, else empty string.",
  "effective_date": "string YYYY-MM-DD - when the agreement takes effect, else null if not stated.",
  "term_months": "integer or null - the agreement's duration/tenure IN MONTHS (convert years to months, e.g. '2 years' -> 24, '1 year' -> 12). Null if no fixed term is stated (e.g. open-ended/perpetual).",
  "end_date": "string YYYY-MM-DD or null - ONLY if the document states an explicit end date outright. Leave null if the end date would need to be calculated from effective_date + term - that calculation is done in code, not by you.",
  "auto_renews": "boolean - true if the agreement states it automatically renews/extends unless cancelled, else false.",
  "renewal_notice_days": "integer or null - how many days of notice are required to prevent auto-renewal or to terminate, else null if not stated.",
  "fee_amount": "number or null - the recurring or total fee amount as a plain number (no currency symbols/commas), else null if not stated.",
  "fee_currency": "string - e.g. 'USD', 'EUR', 'INR' - empty string if unclear or not applicable.",
  "fee_frequency": "string - one of: 'monthly', 'quarterly', 'annual', 'one-time', 'other', or empty string if no fee is stated.",
  "payment_terms": "string - a short free-text summary of payment terms (e.g. 'Net 30, invoiced quarterly in advance'), empty string if none stated.",
  "termination_terms": "string - a short free-text summary of how/when either party can terminate (notice period, cause requirements), empty string if none stated.",
  "other_terms": "string - a short summary of other notable clauses not covered above (confidentiality, exclusivity, SLAs, governing law/jurisdiction, liability caps, etc). Empty string if nothing else notable.",
  "signatories": [
    {
      "name": "string - person's full name",
      "role": "string - their title/role as stated (e.g. 'Chief Executive Officer'), empty string if not given",
      "side": "string - 'client' if they represent the client company, 'provider' if they represent the provider/vendor. Best guess from context if not explicitly labeled."
    }
  ]
}

Be faithful to the document - do not invent figures, dates, or names that aren't stated or clearly
implied. If a field genuinely isn't present in the document, use null/empty string/empty list as
specified above rather than guessing."""


def structure_agreement(raw_text: str) -> dict:
    client = get_client()
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": raw_text},
        ],
        temperature=0.1,
        response_format={"type": "json_object"},
    )
    content = response.choices[0].message.content
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Agreement structuring did not return valid JSON. Raw output:\n{content}") from e


def compute_end_date(effective_date: Optional[str], term_months: Optional[int],
                      explicit_end_date: Optional[str]) -> Optional[str]:
    """Deterministic date math in code rather than asking the LLM to do it
    (see date_utils.py's resolve_relative_phrase for the earlier lesson
    this follows - LLMs are unreliable at exact calendar arithmetic).
    Prefers an explicit end date the document actually states; otherwise
    computes effective_date + term_months when both are known."""
    if explicit_end_date:
        return explicit_end_date
    if effective_date and term_months:
        try:
            start = date.fromisoformat(effective_date)
        except ValueError:
            return None
        return (start + relativedelta(months=term_months)).isoformat()
    return None


def extract_agreement_info(file_bytes: bytes, filename: str) -> dict:
    """Full pipeline: extract raw text from the document (any supported
    format), structure it via the LLM, then fill in a computed end_date
    when the document didn't state one outright. Raises ValueError on
    unreadable/unsupported documents or malformed LLM output."""
    raw_text = extract_text(file_bytes, filename)
    fields = structure_agreement(raw_text)
    fields["end_date"] = compute_end_date(
        fields.get("effective_date"), fields.get("term_months"), fields.get("end_date")
    )
    return fields


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python document_extract.py <path_to_document>")
    else:
        path = sys.argv[1]
        with open(path, "rb") as f:
            result = extract_agreement_info(f.read(), path.split("/")[-1])
        print(json.dumps(result, indent=2))
